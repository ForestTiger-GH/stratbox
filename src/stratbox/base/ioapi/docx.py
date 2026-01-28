"""
docx — чтение/запись DOCX поверх FileStore.

Важно:
- Формат .doc (старый бинарный Word) намеренно не разбирается: его лучше хранить как bytes.
- Для .docx используется python-docx (обычно уже установлен в окружении).

Основные сценарии:
- read_text(): получить весь текст документа одной строкой
- write_text(): быстро сформировать простой DOCX из текста (без сложного форматирования)
"""

from __future__ import annotations

from io import BytesIO
from typing import Iterable

from stratbox.base.filestore.base import FileStore
from stratbox.base.ioapi.bytes import read_bytes, write_bytes


def read_text(path: str, store: FileStore | None = None) -> str:
    """Читает DOCX и возвращает весь текст документа."""
    try:
        import docx  # python-docx
    except Exception as e:
        raise ImportError(
            "DOCX support requires optional dependency 'python-docx'. "
            "Install: pip install python-docx"
        ) from e

    data = read_bytes(path, store=store)
    bio = BytesIO(data)
    d = docx.Document(bio)
    parts: list[str] = []
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)
    return "\n".join(parts)


def write_text(path: str, text: str | Iterable[str], store: FileStore | None = None) -> None:
    """
    Пишет простой DOCX, где каждая строка — отдельный абзац.

    text:
      - str: многострочный текст
      - Iterable[str]: набор строк/абзацев
    """
    try:
        import docx  # python-docx
    except Exception as e:
        raise ImportError(
            "DOCX support requires optional dependency 'python-docx'. "
            "Install: pip install python-docx"
        ) from e

    d = docx.Document()

    if isinstance(text, str):
        lines = text.splitlines()
    else:
        lines = list(text)

    for line in lines:
        d.add_paragraph(str(line))

    out = BytesIO()
    d.save(out)
    write_bytes(path, out.getvalue(), store=store)
